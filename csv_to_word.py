from time import sleep
import sys
import csv
from docx import Document


from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml



RANGE = range(3, 8)

class FailureMode:
	#def __init__(self, number, fmIndex, fm, hazardID, localEffect, finalEffect, ddMethod, currentControls, severity, confidence, causes):
	def __init__(self, number, desc, fmIndex, fm, hazardID, localEffect, finalEffect, ddMethod, currControls, severity, confidence, causes):
		self.number = number
		self.desc = desc
		self.fmIndex = fmIndex
		self.fm = fm
		self.hazardID = hazardID
		self.localEffect = localEffect
		self.finalEffect = finalEffect
		self.ddMethod = ddMethod
		self.currControls = currControls
		self.severity = severity
		self.confidence = confidence
		self.causes = causes

	def addCause(self, cause):
		self.causes += [cause]

	def computeHeight(self):
		numOfCauses = len(self.causes)
		return 10 + numOfCauses * 4

class Causes:
	def __init__(self, number, cause, occurenceScale, improvements, verifyingControl):
		self.number = number
		self.cause = cause
		self.occurenceScale = occurenceScale
		self.improvements = improvements
		self.verifyingControl = verifyingControl


def parseData(data):

	modes = []
	for row in data:
		if any(row[0]):
			cause = Causes(row[11], row[12], row[13], row[14], row[15])
			failuremode = FailureMode(row[0], row[1], row[2], row[3], row[4], row[5], row[6], row[7], row[8], row[9], row[10], [cause])
			modes += [failuremode]
		else:
			cause = Causes(row[11], row[12], row[13], row[14], row[15])
			modes[-1].addCause(cause)
	return modes

def set_column_width(column, width):
	column.width = width 
	for cell in column.cells:
		cell.width = width


def word(filename, output):
    f = open(filename, 'r')
    csv_f = csv.reader(f)
    data = [i for i in list(csv_f) if any(i)]

    headers = data[0]
    data = data[1:]
    numOfModes = 0
    rows = 0
    cols = 5

    for i in data:
    	if i[0]:
    		numOfModes += 1
    		rows += 14
    	else:
    		rows += 4 # add additional possible failure causes

    document = Document('template.docx')
    

    # print '---------'
    
    # for i in range(len(data[0])):
    # 	print str(i) + ' ' + data[0][i]


    parsedData = parseData(data);
    table = document.add_table(rows, cols)
    table.style = 'Style1'
    table.autofit = False
    set_column_width(table.columns[0], 349000)
    set_column_width(table.columns[1], 993000)
    set_column_width(table.columns[2], 349000)
    set_column_width(table.columns[3], 993000)
    set_column_width(table.columns[4], 3493000)
    currOrigin = [0,0]

    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))

    for modeNum in range(len(parsedData)):
    	currMode = parsedData[modeNum]
    	height = currMode.computeHeight() + 1


    	table.cell(currOrigin[0], 0).text = str(currMode.number)
    	table.cell(currOrigin[0], 0).merge(table.cell(currOrigin[0]+height-2,0))
    	if modeNum % 2 != 0:
    		shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    		table.cell(currOrigin[0],0)._tc.get_or_add_tcPr().append(shading_elm)

    	for modeRow in range(0,11):
    		table.cell(currOrigin[0]+modeRow, 1).text = headers[modeRow+1]

    	preCauseAttributes = [currMode.desc, currMode.fmIndex, currMode.fm, currMode.hazardID, currMode.localEffect, currMode.finalEffect, currMode.ddMethod, currMode.currControls, currMode.severity, currMode.confidence]
    	for i in range(len(preCauseAttributes)):
    		table.cell(currOrigin[0]+i, 2).text = preCauseAttributes[i]
    		table.cell(currOrigin[0]+i, 2).merge(table.cell(currOrigin[0]+i, 4))

    	numCauses = len(currMode.causes)
    	table.cell(currOrigin[0]+10,1).merge(table.cell(currOrigin[0]+10+numCauses*4-1,1))

    	for causeNum in range(len(currMode.causes)):
    		causeInst = [currMode.causes[causeNum].cause, currMode.causes[causeNum].occurenceScale, currMode.causes[causeNum].improvements, currMode.causes[causeNum].verifyingControl]
    		offset = 10 + (causeNum)*4
    		table.cell(currOrigin[0]+offset,2).text = str(causeNum+1)
    		if causeNum % 2 != 0:
    			shading_elm_light = parse_xml(r'<w:shd {} w:fill="E5E5E5"/>'.format(nsdecls('w')))
    			table.cell(currOrigin[0]+offset,2)._tc.get_or_add_tcPr().append(shading_elm_light)
    		table.cell(currOrigin[0]+offset,2).merge(table.cell(currOrigin[0]+offset+3,2))
    		for causeRow in range(0,4):
    			table.cell(currOrigin[0]+offset+causeRow,3).text = headers[causeRow+12]
    		table.cell(currOrigin[0]+offset+0,4).text = causeInst[0]
    		table.cell(currOrigin[0]+offset+1,4).text = causeInst[1]
    		causeInst[2] = [x.strip() for x in causeInst[2].split('~') if x != '']
    		filter(None, causeInst[2])
    		if any(causeInst[2]):
	    		table.cell(currOrigin[0]+offset+2,4).text = causeInst[2][0]
	    		table.cell(currOrigin[0]+offset+2,4).paragraphs[0].style = 'Bulletz'
    		for i in causeInst[2][1:]:
    			table.cell(currOrigin[0]+offset+2,4).add_paragraph(i,'Bulletz') 

    		causeInst[3] = [x.strip() for x in causeInst[3].split('~') if x != '']
    		filter(None, causeInst[3])
    		if any(causeInst[3]):
    			table.cell(currOrigin[0]+offset+3,4).text = causeInst[3][0]
    			table.cell(currOrigin[0]+offset+3,4).paragraphs[0].style = 'Bulletz'
    		for i in causeInst[3][1:]:
    			table.cell(currOrigin[0]+offset+3,4).add_paragraph(i,'Bulletz')


    	
    	currOrigin[0] += height-1



    		#remember to merge cells 1st col before moving onto next ModeNum

   	# Set a cell background (shading) color to RGB D9D9D9. 
	#shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
	#table.cell(0,0)._tc.get_or_add_tcPr().append(shading_elm)

    document.save(output)
   	
 
if __name__ == '__main__':
    word(*sys.argv[1:])
